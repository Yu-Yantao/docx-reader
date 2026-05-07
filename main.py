"""读取 Word 文档并还原段落的自动编号文本。

docx 把"编号格式"和"编号序号"分开存放：
- 编号格式（numFmt、lvlText、start、suff 等）集中定义在 word/numbering.xml 里，
  描述的是"这套列表第 N 级长什么样"。
- 段落自身（word/document.xml）只通过 <w:numPr> 记录 numId + ilvl 这一对引用，
  说明"我属于哪套列表、第几级"。
- 真正显示的"1."、"一、"、"(A)" 既不在格式定义里，也不在段落里，而是 Word 在打开
  文档时按段落出现顺序边走边累加计数后渲染出来的。

所以直接用 python-docx 读 paragraph.text，只能拿到段落正文，看不到自动编号。
本模块把 numbering.xml 解析出来，再按段落顺序自己维护一份计数，把编号文本还原出来
拼回段落前面。
"""

import logging
import re
from typing import List

from docx import Document
from docx.oxml.ns import qn


class FakeParagraph:
    """轻量段落对象，承载拼接编号后的解析结果。"""

    def __init__(self, text, style, runs):
        self.text = text
        self.style = style
        self.runs = runs


class DocxReader:
    """解析 docx numbering 定义，遍历段落时动态计数并还原编号前缀。"""

    ideographTraditional = "甲乙丙丁戊己庚辛壬癸"
    ideographZodiac = "子丑寅卯辰巳午未申酉戌亥"

    def __init__(self, docx, gap_text="\t"):
        self.docx = Document(docx)
        self.numId2style = self.get_style_data()
        self.gap_text = gap_text
        # (numId, ilvl) -> 当前计数值（int）
        self.cnt = {}
        # (numId, ilvl) -> 格式化后的编号文本，供多级模板 "%1.%2" 引用
        self.cache = {}
        self.result = []

    @property
    def paragraphs(self) -> List[FakeParagraph]:
        """返回补齐编号文本后的段落列表，首次访问时解析，后续返回缓存。"""
        if self.result:
            return self.result.copy()
        self.cnt.clear()
        self.cache.clear()
        for paragraph in self.docx.paragraphs:
            pPr = paragraph._element.pPr
            numPr = pPr.numPr if pPr is not None else None
            number_text = self.get_number_text(numPr)
            # 段落内手动换行只给第一行挂编号
            parts = paragraph.text.split("\n")
            for i, part in enumerate(parts):
                fake_p = FakeParagraph(part, paragraph.style, paragraph.runs)
                if i == 0:
                    fake_p.text = number_text + fake_p.text
                self.result.append(fake_p)
        return self.result.copy()

    def get_style_data(self):
        """解析 numbering.xml，建立 (numId, ilvl) -> 样式属性 的映射。

        numbering.xml 里有两层结构：
        - <w:abstractNum> 是真正的格式模板，定义每一级的 numFmt、lvlText、start、suff 等。
        - <w:num> 是模板的实例化，把 abstractNumId 绑到一个 numId 上，段落里引用的是 numId。
        这种分层是为了同一份模板能被多个 numId 复用、各自独立计数。
        这里把两层关系打平成 (numId, ilvl) -> style 的扁平表，后续查样式直接命中。
        """
        # 检查文档是否包含 numbering 部分
        try:
            self.docx.part.numbering_part._element
        except NotImplementedError:
            return {}

        numbering_part = self.docx.part.numbering_part._element

        # abstractNumId -> numId 的映射
        abstractId2numId = {num.abstractNumId.val: num.numId for num in numbering_part.num_lst}

        numId2style = {}

        for abstractNumIdTag in numbering_part.findall(qn("w:abstractNum")):
            abstractNumId = abstractNumIdTag.get(qn("w:abstractNumId"))
            numId = abstractId2numId[int(abstractNumId)]

            for lvlTag in abstractNumIdTag.findall(qn("w:lvl")):
                ilvl = lvlTag.get(qn("w:ilvl"))

                # 提取该级别所有带 w:val 属性的标签作为样式属性
                style = {
                    tag.tag[tag.tag.rfind("}") + 1:]: tag.get(qn("w:val")) for tag in
                    lvlTag.xpath("./*[@w:val]", namespaces=numbering_part.nsmap)
                }

                # 兼容处理：numFmt 可能藏在 AlternateContent 里
                if "numFmt" not in style:
                    numFmtVal = lvlTag.xpath("./mc:AlternateContent/mc:Fallback/w:numFmt/@w:val",
                                             namespaces=numbering_part.nsmap)
                    if numFmtVal and numFmtVal[0] == "decimal":
                        numFmt_format = lvlTag.xpath("./mc:AlternateContent/mc:Choice/w:numFmt/@w:format",
                                                     namespaces=numbering_part.nsmap)
                        if numFmt_format:
                            style["numFmt"] = "decimal" + numFmt_format[0].split(",")[0]

                if style.get("numFmt") == "decimalZero":
                    style["numFmt"] = "decimal01"

                numId2style[(numId, int(ilvl))] = style

        return numId2style

    @staticmethod
    def int2upperLetter(num):
        """1->A, 2->B, 27->AA"""
        result = []
        while num > 0:
            num -= 1
            remainder = num % 26
            result.append(chr(remainder + ord("A")))
            num //= 26
        return "".join(reversed(result))

    @staticmethod
    def int2upperRoman(num):
        """整数转大写罗马数字。"""
        t = [
            (1000, "M"),
            (900, "CM"),
            (500, "D"),
            (400, "CD"),
            (100, "C"),
            (90, "XC"),
            (50, "L"),
            (40, "XL"),
            (10, "X"),
            (9, "IX"),
            (5, "V"),
            (4, "IV"),
            (1, "I"),
        ]
        roman_num = ""
        i = 0
        while num > 0:
            val, syb = t[i]
            for _ in range(num // val):
                roman_num += syb
                num -= val
            i += 1
        return roman_num

    @staticmethod
    def int2cardinalText(num):
        """整数转英文基数词：21 -> Twenty-one"""
        if not isinstance(num, int) or num < 0 or num > 999999999:
            raise ValueError("Invalid number: must be a positive integer within four digits")
        base = [
            "Zero",
            "One",
            "Two",
            "Three",
            "Four",
            "Five",
            "Six",
            "Seven",
            "Eight",
            "Nine",
            "Ten",
            "Eleven",
            "Twelve",
            "Thirteen",
            "Fourteen",
            "Fifteen",
            "Sixteen",
            "Seventeen",
            "Eighteen",
            "Nineteen",
        ]
        tens = ["", "", "Twenty", "Thirty", "Fourty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]
        thousands = ["", "Thousand", "Million", "Billion"]

        def two_digits(n):
            if n < 20:
                return base[n]
            ten, unit = divmod(n, 10)
            if unit == 0:
                return f"{tens[ten]}"
            else:
                return f"{tens[ten]}-{base[unit]}"

        def three_digits(n):
            hundred, rest = divmod(n, 100)
            if hundred == 0:
                return two_digits(rest)
            result = f"{base[hundred]} hundred "
            if rest > 0:
                result += two_digits(rest)
            return result.strip()

        if num < 99:
            return two_digits(num)
        chunks = []
        while num > 0:
            num, remainder = divmod(num, 1000)
            chunks.append(remainder)
        words = []
        for i in range(len(chunks) - 1, -1, -1):
            if chunks[i] == 0:
                continue
            chunk_word = three_digits(chunks[i])
            if thousands[i]:
                chunk_word += f" {thousands[i]}"
            words.append(chunk_word)
        words = " ".join(words).lower()
        return words[0].upper() + words[1:]

    @staticmethod
    def int2ordinalText(num):
        """整数转英文序数词：21 -> Twenty-first"""
        if not isinstance(num, int) or num < 0 or num > 999999:
            raise ValueError("Invalid number: must be a positive integer within four digits")
        base = [
            "Zero",
            "One",
            "Two",
            "Three",
            "Four",
            "Five",
            "Six",
            "Seven",
            "Eight",
            "Nine",
            "Ten",
            "Eleven",
            "Twelve",
            "Thirteen",
            "Fourteen",
            "Fifteen",
            "Sixteen",
            "Seventeen",
            "Eighteen",
            "Nineteen",
        ]
        baseth = [
            "Zeroth",
            "First",
            "Second",
            "Third",
            "Fourth",
            "Fifth",
            "Sixth",
            "Seventh",
            "Eighth",
            "Ninth",
            "Tenth",
            "Eleventh",
            "Twelfth",
            "Thirteenth",
            "Fourteenth",
            "Fifteenth",
            "Sixteenth",
            "Seventeenth",
            "Eighteenth",
            "Nineteenth",
            "Twentieth",
        ]
        tens = ["", "", "Twenty", "Thirty", "Fourty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]
        tensth = ["", "", "Twentieth", "Thirtieth", "Fortieth", "Fiftieth", "Sixtieth", "Seventieth", "Eightieth",
                  "Ninetieth"]

        def two_digits(n):
            if n <= 20:
                return baseth[n]
            ten, unit = divmod(n, 10)
            result = tensth[ten]
            if unit != 0:
                result = f"{tens[ten]}-{baseth[unit]}"
            return result

        thousand, num = divmod(num, 1000)
        result = []
        if thousand > 0:
            if num == 0:
                return f"{DocxReader.int2cardinalText(thousand)} thousandth"
            result.append(f"{DocxReader.int2cardinalText(thousand)} thousand")
        hundred, num = divmod(num, 100)
        if hundred > 0:
            if num == 0:
                result.append(f"{base[hundred]} hundredth")
                return " ".join(result)
            result.append(f"{base[hundred]} hundred")
        result.append(two_digits(num))
        result = " ".join(result).lower()
        return result[0].upper() + result[1:]

    @staticmethod
    def int2Chinese(num, ch_num, units):
        """按给定数字表和单位表转中文数字。"""
        if not (0 <= num <= 99999999):
            raise ValueError("仅支持小于一亿以内的正整数")

        def int2Chinese_in(num, ch_num, units):
            """处理 0~9999，外层负责拼接"万"。"""
            if not (0 <= num <= 9999):
                raise ValueError("仅支持小于一万以内的正整数")
            result = [ch_num[int(i)] + unit for i, unit in zip(reversed(str(num).zfill(4)), units)]
            result = "".join(reversed(result))
            zero_char = ch_num[0]
            result = re.sub(f"(?:{zero_char}[{units}])+", zero_char, result)
            result = result.rstrip(units[0])
            if result != zero_char:
                result = result.rstrip(zero_char)
            if result.lstrip(zero_char).startswith("一十"):
                result = result.replace("一", "")
            return result

        if num < 10000:
            result = int2Chinese_in(num, ch_num, units)
        else:
            left = num // 10000
            right = num % 10000
            result = int2Chinese_in(left, ch_num, units) + "万" + int2Chinese_in(right, ch_num, units)
        if result != ch_num[0]:
            result = result.strip(ch_num[0])
        return result

    @staticmethod
    def int2ChineseCounting(num):
        """12 -> 十二"""
        return DocxReader.int2Chinese(num, ch_num="〇一二三四五六七八九", units="个十百千")

    @staticmethod
    def int2ChineseLegalSimplified(num):
        """12 -> 拾贰"""
        return DocxReader.int2Chinese(num, ch_num="零壹贰叁肆伍陆柒捌玖", units="个拾佰仟")

    def get_number_text(self, numpr):
        """根据段落的 numPr 引用，结合样式定义和当前计数，还原编号前缀。

        Word 渲染时是按段落顺序边走边累加的，这里复刻同样的过程：
        每碰到一个挂着 (numId, ilvl) 的段落，就把对应层级的计数 +1，
        再用样式里的 numFmt 把整数转成"一"、"A"、"01" 等显示形式，
        最后套进 lvlText 模板（多级编号会引用上层缓存的文本）。
        """
        if numpr is None or numpr.numId.val == 0:
            return ""
        numId = numpr.numId.val
        ilvl = numpr.ilvl.val

        style = self.numId2style[(numId, ilvl)]
        numFmt: str = style.get("numFmt")
        lvlText = style.get("lvlText")

        # 当前层级第一次出现时从 start 起算，之后每命中一次就 +1
        if (numId, ilvl) in self.cnt:
            self.cnt[(numId, ilvl)] += 1
        else:
            self.cnt[(numId, ilvl)] = int(style["start"])
        pos = self.cnt[(numId, ilvl)]
        num_text = str(pos)

        # 按 numFmt 把整数序号转成对应的显示格式
        if numFmt.startswith("decimal"):
            num_text = num_text.zfill(numFmt.count("0") + 1)
        elif numFmt == "upperRoman":
            num_text = self.int2upperRoman(pos)
        elif numFmt == "lowerRoman":
            num_text = self.int2upperRoman(pos).lower()
        elif numFmt == "upperLetter":
            num_text = self.int2upperLetter(pos)
        elif numFmt == "lowerLetter":
            num_text = self.int2upperLetter(pos).lower()
        elif numFmt == "ordinal":
            num_text = f"{pos}{'th' if 11 <= pos <= 13 else {1: 'st', 2: 'nd', 3: 'rd'}.get(pos % 10, 'th')}"
        elif numFmt == "cardinalText":
            num_text = self.int2cardinalText(pos)
        elif numFmt == "ordinalText":
            num_text = self.int2ordinalText(pos)
        elif numFmt == "ideographTraditional":
            if 1 <= pos <= 10:
                num_text = self.ideographTraditional[pos - 1]
        elif numFmt == "ideographZodiac":
            if 1 <= pos <= 12:
                num_text = self.ideographZodiac[pos - 1]
        elif numFmt == "chineseCounting":
            num_text = self.int2ChineseCounting(pos)
        elif numFmt == "chineseLegalSimplified":
            num_text = self.int2ChineseLegalSimplified(pos)
        elif numFmt == "decimalEnclosedCircleChinese":
            pass

        # 缓存当前层级文本，多级模板如 "%1.%2" 会引用各层级的缓存
        self.cache[(numId, ilvl)] = num_text
        for i in range(0, ilvl + 1):
            lvlText = lvlText.replace(f"%{i + 1}", self.cache.get((numId, i), ""))

        suff_text = {"space": " ", "nothing": ""}.get(style.get("suff"), self.gap_text)
        lvlText += suff_text
        return lvlText


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO, format="%(asctime)-24s  %(levelname)-6s  %(filename)12s:%(lineno)-4d  %(message)s")

    def build_preview(lines, prefix_len=15):
        content = []
        for index, line in enumerate(lines, start=1):
            content.append(f"{str(line)[:prefix_len]}")
        return "\n".join(content)

    docx_path = r"test.docx"

    # python-docx 直接读：拿不到自动编号
    raw_doc = Document(docx_path)
    raw_content = build_preview(paragraph.text for paragraph in raw_doc.paragraphs)

    # DocxReader 读：还原编号
    parsed_doc = DocxReader(docx_path, "")
    parsed_content = build_preview(paragraph.text for paragraph in parsed_doc.paragraphs)

    print("=== python-docx 直接读取（不含自动编号） ===")
    print(raw_content)
    print()
    print("=== DocxReader 读取（还原自动编号） ===")
    print(parsed_content)
